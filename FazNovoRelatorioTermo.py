from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from python_docx_replace import docx_replace
from datetime import datetime
import pandas as pd

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

tabela = pd.read_excel("TabelaValoresTemperatura.xlsx", sheet_name="InfoTermo")

documento = Document("RelatórioTermográficoVariaveis.docx")


for linha in tabela.index:

    localtit = str(tabela.loc[linha, "Título"])
    mestermo = str(tabela.loc[linha, "MesTermo"])
    anotermo = str(tabela.loc[linha, "AnoTermo"])
    localac = str(tabela.loc[linha,"Local A/C"])
    endereco = str(tabela.loc[linha, "Endereço"])
    diasana = str(tabela.loc[linha, "Dia(s) da análise"])
    maxtempmod = str(tabela.loc[linha, "Maior temperatura de funcionamento do módulo"])
    quanttermo = str(tabela.loc[linha, "Quantidade de termogramas"])
    horini = str(tabela.loc[linha, "Dia, Hora Início"])
    horter = str(tabela.loc[linha, "Dia, Hora Término"])
    maxtempamb = str(tabela.loc[linha, "Temperatura ambiente máxima"])
    modelomod = str(tabela.loc[linha, "Modelo do módulo"])
    tempmed = str(tabela.loc[linha, "Valor médio das medições"])
    porabaixo = str(tabela.loc[linha, "Porcentagem abaixo da temperatura limite"])

    tabela3 = pd.read_excel("TabelaValoresTemperatura.xlsx", sheet_name="DadosTabela3")
    texto = " "
    for k in tabela3.index:
        if tabela3.iloc[k, 3] != 0:
            texto = texto + str(tabela3.iloc[k, 4]) + ", "
    novotexto = texto[:-2] + " e nenhum apresentou temperatura maior que " + texto[-6:-2] + "."

    my_dict = {
        "TITULO": localtit,
        "MESTERMO": mestermo,
        "ANOTERMO": anotermo,
        "LOCALCLIENTE": localac,
        "ENDERECOLOCAL": endereco,
        "DIASANALISE": diasana,
        "MAXTEMPMOD": maxtempmod,
        "QUANTTERMO": quanttermo,
        "HORINI": horini,
        "HORTER": horter,
        "MAXTEMPAMB": maxtempamb,
        "MODELOMOD": modelomod,
        "TEMPMED": tempmed,
        "PORABAIXO": porabaixo,
        "TEXTOINTERV": novotexto,
    }

    docx_replace(documento, **my_dict)

    # tabela2 = pd.read_excel("C://Users/ORION/Desktop/Automacao/Tabela2.xlsx")
    # for paragrafo in documento.paragraphs:
    #     if "Tabela 2" in paragrafo.text:
    #         p = paragrafo.insert_paragraph_before('INSIRA A TABELA AQUI')
            # table = documento.add_table(rows= 1, cols=4)
            # k=0;
            # for col in tabela2.columns:
            #     table.cell(0,k).text = col
            #     k = k+1;
            #
            # for i in range(len(tabela2)):
            #     cells = table.add_row().cells
            #     cells[0].text = str(tabela2.iloc[i,0])
            #     cells[1].text = str(tabela2.iloc[i,1])
            #     cells[2].text = str(tabela2.iloc[i,2])
            #     cells[3].text = str(tabela2.iloc[i,3])
            #
            # tabela2.style = 'TabeladeGrade4'
            #
            # move_table_after(table,p)

    # k=1
    # for paragrafo in documento.paragraphs:
    #     if 'NÚMERO DE OCORRÊNCIAS DAS TEMPERATURAS' in paragrafo.text:
    #         p = paragrafo.insert_paragraph_before()
    #         run = p.add_run()
    #         run.add_picture('C://Users/ORION/Desktop/Automacao/Gráfico1.png', height=Cm(7.5))




    documento.save("RelatorioTermo.docx")